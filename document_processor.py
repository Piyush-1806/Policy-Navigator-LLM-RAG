import asyncio
from email import parser
import tempfile
import os
from typing import List, Dict, Any
from urllib.parse import urlparse
import logging
import email
from email import policy

import httpx
import fitz  # PyMuPDF
from docx import Document
import nltk
from sentence_transformers import SentenceTransformer

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.embedding_model = None
        
    async def initialize(self):
        """Initialize the document processor"""
        try:
            # Download NLTK data with specific resources
            logger.info("Downloading NLTK resources...")
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('punkt_tab', quiet=True)
                nltk.download('stopwords', quiet=True)
                logger.info("NLTK data downloaded successfully")
            except Exception as e:
                logger.warning(f"Some NLTK downloads failed: {e}")
        except Exception as e:
            logger.warning(f"NLTK initialization failed: {e}")
        
        # Load embedding model
        try:
            self.embedding_model = SentenceTransformer('all-mpnet-base-v2')
            logger.info("Embedding model loaded successfully")
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
        
    async def process_document(self, url: str) -> List[Dict[str, Any]]:
        """Download and process a document into chunks"""
        try:
            # Download document
            content = await self._download_document(url)
            
            # Parse based on file type
            text = await self._parse_document(url, content)
            
            # Validate the extracted text
            if self._is_text_garbage(text):
                logger.error("Extracted text appears to be garbage/corrupted")
                raise ValueError("Document parsing failed - extracted text is corrupted")
            
            # Create chunks
            chunks = self._create_chunks(text, url)
            
            logger.info(f"Successfully processed document: {len(chunks)} chunks created")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing document {url}: {e}")
            raise
    
    def _is_text_garbage(self, text: str) -> bool:
        """Check if extracted text is garbage/corrupted"""
        if not text or len(text.strip()) < 100:
            return True
        
        # Check for excessive non-ASCII characters
        non_ascii_count = sum(1 for char in text if ord(char) > 127)
        non_ascii_ratio = non_ascii_count / len(text) if text else 1
        
        if non_ascii_ratio > 0.3:  # More than 30% non-ASCII
            logger.warning(f"High non-ASCII character ratio: {non_ascii_ratio:.2f}")
            return True
        
        # Check for excessive special characters
        special_chars = sum(1 for char in text if not char.isalnum() and char not in ' \n\t.,!?;:()')
        special_ratio = special_chars / len(text) if text else 1
        
        if special_ratio > 0.2:  # More than 20% special characters
            logger.warning(f"High special character ratio: {special_ratio:.2f}")
            return True
        
        # Check if we have actual words
        words = text.split()
        if len(words) < 50:  # Too few words
            return True
        
        # Check for reasonable word lengths
        avg_word_length = sum(len(word) for word in words[:100]) / min(len(words), 100)
        if avg_word_length > 15 or avg_word_length < 2:  # Unreasonable word lengths
            logger.warning(f"Unreasonable average word length: {avg_word_length}")
            return True
        
        return False
    
    async def _download_document(self, url: str) -> bytes:
        """Download document from URL"""
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url)
                response.raise_for_status()
                logger.info(f"Downloaded document: {len(response.content)} bytes")
                return response.content
        except Exception as e:
            logger.error(f"Failed to download document from {url}: {e}")
            raise
    
    async def _parse_document(self, url: str, content: bytes) -> str:
        """Parse document content with multiple parsing strategies"""
        file_ext = os.path.splitext(urlparse(url).path)[1].lower()
        
        # Also check Content-Type if extension is unclear
        content_type = self._detect_content_type(content)
        
        if file_ext == '.pdf' or 'pdf' in content_type:
            return await self._parse_pdf_robust(content)
        elif file_ext in ['.docx', '.doc'] or 'officedocument' in content_type or 'msword' in content_type:
            return await self._parse_docx(content)
        elif file_ext in ['.eml', '.msg'] or content_type.startswith('message/'):
            return await self._parse_email(content)
        elif file_ext in ['.txt', '.md']:
            return await self._parse_text(content)
        else:
            # Try to detect format by content
            return await self._parse_unknown_format(content, url)
    
    def _detect_content_type(self, content: bytes) -> str:
        """Detect content type from file signature"""
        if content.startswith(b'%PDF'):
            return 'application/pdf'
        elif content.startswith(b'PK') and b'word/' in content[:1000]:
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif content.startswith(b'\xd0\xcf\x11\xe0'):  # Old MS Office format
            return 'application/msword'
        elif content.startswith(b'From ') or b'Message-ID:' in content[:1000]:
            return 'message/rfc822'
        else:
            return 'text/plain'
    
    async def _parse_pdf_robust(self, content: bytes) -> str:
        """Robust PDF parsing with multiple strategies"""
        temp_file = None
        
        try:
            # Strategy 1: Standard PyMuPDF extraction
            temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            temp_file.write(content)
            temp_file.flush()
            temp_file.close()
            
            doc = fitz.open(temp_file.name)
            text = ""
            
            logger.info(f"PDF has {doc.page_count} pages")
            
            for page_num in range(doc.page_count):
                try:
                    page = doc[page_num]
                    
                    # Try multiple extraction methods
                    page_text = ""
                    
                    # Method 1: Standard text extraction
                    try:
                        page_text = page.get_text()
                        if page_text.strip():
                            logger.debug(f"Page {page_num + 1}: Standard extraction successful")
                    except Exception as e:
                        logger.warning(f"Page {page_num + 1}: Standard extraction failed: {e}")
                    
                    # Method 2: Try different text extraction flags if standard fails
                    if not page_text.strip():
                        try:
                            page_text = page.get_text("text")
                            if page_text.strip():
                                logger.debug(f"Page {page_num + 1}: Text flag extraction successful")
                        except Exception as e:
                            logger.warning(f"Page {page_num + 1}: Text flag extraction failed: {e}")
                    
                    # Method 3: Try blocks-based extraction
                    if not page_text.strip():
                        try:
                            blocks = page.get_text("dict")["blocks"]
                            block_text = ""
                            for block in blocks:
                                if "lines" in block:
                                    for line in block["lines"]:
                                        for span in line["spans"]:
                                            block_text += span["text"] + " "
                            page_text = block_text
                            if page_text.strip():
                                logger.debug(f"Page {page_num + 1}: Block extraction successful")
                        except Exception as e:
                            logger.warning(f"Page {page_num + 1}: Block extraction failed: {e}")
                    
                    # Add page text if we got something meaningful
                    if page_text.strip() and len(page_text.strip()) > 10:
                        text += f"[Page {page_num + 1}] {page_text}\n"
                    else:
                        logger.warning(f"Page {page_num + 1}: No meaningful text extracted")
                        
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1}: {e}")
                    continue
            
            doc.close()
            
            if text.strip():
                logger.info(f"PDF parsed successfully: {len(text)} characters extracted")
                return text
            else:
                logger.error("No text could be extracted from PDF")
                raise ValueError("PDF text extraction failed - no readable content found")
                
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            
            # Fallback: Try to extract text using different approach
            try:
                logger.info("Trying fallback PDF parsing...")
                doc = fitz.open(stream=content, filetype="pdf")
                fallback_text = ""
                
                for page_num in range(min(5, doc.page_count)):  # Try first 5 pages
                    try:
                        page = doc[page_num]
                        page_text = page.get_text()
                        if page_text and len(page_text.strip()) > 20:
                            fallback_text += f"[Page {page_num + 1}] {page_text}\n"
                    except:
                        continue
                
                doc.close()
                
                if fallback_text.strip():
                    logger.info("Fallback PDF parsing successful")
                    return fallback_text
                    
            except Exception as fallback_error:
                logger.error(f"Fallback PDF parsing also failed: {fallback_error}")
            
            raise ValueError("All PDF parsing methods failed")
            
        finally:
            # Clean up temp file
            if temp_file and os.path.exists(temp_file.name):
                try:
                    os.unlink(temp_file.name)
                except:
                    pass
    
    async def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX content with robust error handling"""
        temp_file = None
        try:
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_file.write(content)
            temp_file.flush()
            temp_file.close()
            
            logger.info(f"Created temp DOCX file: {temp_file.name}")
            
            # Try to open with python-docx
            try:
                doc = Document(temp_file.name)
                text = ""
                
                # Extract paragraphs
                paragraph_count = 0
                for paragraph in doc.paragraphs:
                    if paragraph.text and paragraph.text.strip():
                        text += paragraph.text.strip() + "\n"
                        paragraph_count += 1
                
                logger.info(f"Extracted {paragraph_count} paragraphs from DOCX")
                
                # Extract tables if present
                table_count = 0
                for table in doc.tables:
                    table_count += 1
                    text += f"\n[Table {table_count}]\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text and cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text += " | ".join(row_text) + "\n"
                
                if table_count > 0:
                    logger.info(f"Extracted {table_count} tables from DOCX")
                
                # Extract headers and footers if possible
                try:
                    for section in doc.sections:
                        if section.header.paragraphs:
                            header_text = []
                            for para in section.header.paragraphs:
                                if para.text and para.text.strip():
                                    header_text.append(para.text.strip())
                            if header_text:
                                text = "[Header] " + " ".join(header_text) + "\n" + text
                        
                        if section.footer.paragraphs:
                            footer_text = []
                            for para in section.footer.paragraphs:
                                if para.text and para.text.strip():
                                    footer_text.append(para.text.strip())
                            if footer_text:
                                text += "\n[Footer] " + " ".join(footer_text)
                except Exception as e:
                    logger.warning(f"Could not extract headers/footers: {e}")
                
            except Exception as docx_error:
                logger.error(f"python-docx parsing failed: {docx_error}")
                raise ValueError(f"DOCX parsing failed: {docx_error}")
            
            logger.info(f"DOCX parsed successfully: {len(text)} characters")
            
            if not text.strip():
                raise ValueError("DOCX appears to be empty or contains no readable text")
                
            return text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            
            # Try fallback parsing methods
            try:
                logger.info("Attempting fallback DOCX parsing...")
                # Try to extract as ZIP and read document.xml directly
                import zipfile
                import xml.etree.ElementTree as ET
                
                with tempfile.NamedTemporaryFile() as temp_zip:
                    temp_zip.write(content)
                    temp_zip.flush()
                    
                    with zipfile.ZipFile(temp_zip.name, 'r') as zip_ref:
                        # Try to read document.xml
                        try:
                            doc_xml = zip_ref.read('word/document.xml')
                            root = ET.fromstring(doc_xml)
                            
                            # Extract text from XML
                            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                            text_elements = root.findall('.//w:t', namespaces)
                            
                            text = ""
                            for elem in text_elements:
                                if elem.text:
                                    text += elem.text + " "
                            
                            if text.strip():
                                logger.info(f"Fallback DOCX parsing successful: {len(text)} characters")
                                return text
                                
                        except Exception as xml_error:
                            logger.warning(f"XML fallback failed: {xml_error}")
                            
            except Exception as fallback_error:
                logger.error(f"All DOCX fallback methods failed: {fallback_error}")
            
            raise ValueError(f"All DOCX parsing methods failed: {e}")
            
        finally:
            # Clean up temp file
            if temp_file and os.path.exists(temp_file.name):
                try:
                    os.unlink(temp_file.name)
                    logger.debug("Cleaned up temp DOCX file")
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup temp file: {cleanup_error}")
    
    async def _parse_email(self, content: bytes) -> str:
        """Parse email content with multiple strategies"""
        try:
            # Strategy 1: Use Python's built-in email library
            try:
                # Try to parse as email message
                if isinstance(content, bytes):
                    content_str = content.decode('utf-8', errors='ignore')
                else:
                    content_str = content
                
                msg = email.message_from_string(content_str, policy=policy.default)
                
                text = ""
                
                # Extract headers
                if msg.get('Subject'):
                    text += f"Subject: {msg.get('Subject')}\n"
                if msg.get('From'):
                    text += f"From: {msg.get('From')}\n"
                if msg.get('To'):
                    text += f"To: {msg.get('To')}\n"
                if msg.get('Date'):
                    text += f"Date: {msg.get('Date')}\n"
                if msg.get('Cc'):
                    text += f"CC: {msg.get('Cc')}\n"
                
                text += "\n"
                
                # Extract body
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        if content_type == 'text/plain':
                            body = part.get_payload(decode=True)
                            if isinstance(body, bytes):
                                body = body.decode('utf-8', errors='ignore')
                            text += f"Body (Plain Text):\n{body}\n\n"
                        elif content_type == 'text/html':
                            body = part.get_payload(decode=True)
                            if isinstance(body, bytes):
                                body = body.decode('utf-8', errors='ignore')
                            # Simple HTML stripping
                            import re
                            body = re.sub(r'<[^>]+>', '', body)
                            text += f"Body (HTML converted):\n{body}\n\n"
                else:
                    # Single part message
                    body = msg.get_payload(decode=True)
                    if isinstance(body, bytes):
                        body = body.decode('utf-8', errors='ignore')
                    text += f"Body:\n{body}\n"
                
                logger.info(f"Email parsed successfully with email library: {len(text)} characters")
                
                if text.strip():
                    return text
                    
            except Exception as email_error:
                logger.warning(f"Built-in email parsing failed: {email_error}")
            
            # Strategy 2: Try with mailparser if available (optional)
            try:
                # import mail-parser
                mail = mail-parser.parse_from_bytes(content)
                
                text = ""
                if hasattr(mail, 'subject') and mail.subject:
                    text += f"Subject: {mail.subject}\n"
                if hasattr(mail, 'from_') and mail.from_:
                    text += f"From: {mail.from_}\n"
                if hasattr(mail, 'to') and mail.to:
                    text += f"To: {mail.to}\n"
                if hasattr(mail, 'date') and mail.date:
                    text += f"Date: {mail.date}\n"
                if hasattr(mail, 'cc') and mail.cc:
                    text += f"CC: {mail.cc}\n"
                
                text += "\n"
                
                if hasattr(mail, 'body') and mail.body:
                    text += f"Body: {mail.body}\n"
                
                logger.info(f"Email parsed successfully with mailparser: {len(text)} characters")
                
                if text.strip():
                    return text
                    
            except ImportError:
                logger.debug("mailparser not available, using built-in email parser")
            except Exception as mailparser_error:
                logger.warning(f"mailparser failed: {mailparser_error}")
            
            # Strategy 3: Simple text parsing fallback
            logger.info("Using simple text parsing for email")
            text = content.decode('utf-8', errors='ignore')
            
            # Try to clean up common email artifacts
            import re
            
            # Remove quoted-printable encoding artifacts
            text = re.sub(r'=\n', '', text)
            text = re.sub(r'=([A-F0-9]{2})', lambda m: chr(int(m.group(1), 16)), text)
            
            logger.info(f"Email parsed as plain text: {len(text)} characters")
            
            if not text.strip():
                raise ValueError("Email appears to be empty")
                
            return text
            
        except Exception as e:
            logger.error(f"All email parsing methods failed: {e}")
            raise ValueError(f"Email parsing failed: {e}")
    
    async def _parse_text(self, content: bytes) -> str:
        """Parse plain text content"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    text = content.decode(encoding)
                    logger.info(f"Text parsed with {encoding} encoding: {len(text)} characters")
                    return text
                except UnicodeDecodeError:
                    continue
            
            # Fallback with error handling
            text = content.decode('utf-8', errors='ignore')
            logger.info(f"Text parsed with error handling: {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"Text parsing failed: {e}")
            raise ValueError(f"Text parsing failed: {e}")
    
    async def _parse_unknown_format(self, content: bytes, url: str) -> str:
        """Try to parse unknown format by attempting different parsers"""
        logger.info(f"Attempting to detect format for: {url}")
        
        # Try PDF first
        if content.startswith(b'%PDF'):
            logger.info("Detected PDF format")
            return await self._parse_pdf_robust(content)
        
        # Try DOCX
        if content.startswith(b'PK') and b'word/' in content[:1000]:
            logger.info("Detected DOCX format")
            return await self._parse_docx(content)
        
        # Try email
        if b'Message-ID:' in content[:1000] or content.startswith(b'From '):
            logger.info("Detected email format")
            return await self._parse_email(content)
        
        # Default to text
        logger.info("Treating as text format")
        return await self._parse_text(content)
    
    def _create_chunks(self, text: str, source_url: str) -> List[Dict[str, Any]]:
        """Create chunks from text with fallback for NLTK issues"""
        try:
            # Validate text before chunking
            if self._is_text_garbage(text):
                raise ValueError("Text appears to be corrupted")
            
            # Try NLTK sentence tokenization first
            try:
                sentences = nltk.sent_tokenize(text)
                logger.info(f"NLTK tokenization successful: {len(sentences)} sentences")
            except Exception as nltk_error:
                logger.warning(f"NLTK tokenization failed: {nltk_error}")
                sentences = self._simple_sentence_split(text)
                logger.info(f"Fallback sentence splitting: {len(sentences)} sentences")
            
            chunks = []
            chunk_size = 12  # Larger chunks to preserve complete definitions
            overlap = 4
            
            for i in range(0, len(sentences), chunk_size - overlap):
                chunk_sentences = sentences[i:i + chunk_size]
                chunk_text = ' '.join(chunk_sentences)
                
                # Skip very short chunks
                if len(chunk_text.strip()) < 200:  # Increased minimum length for definitions
                    continue
                
                # Special handling for definition chunks - extend if they contain key terms
                if any(term in chunk_text for term in ['means', 'definition', 'defined as', 'refers to']):
                    # Try to extend definition chunks to capture complete definitions
                    extended_sentences = sentences[i:min(i + chunk_size + 4, len(sentences))]
                    extended_text = ' '.join(extended_sentences)
                    if len(extended_text) < 2000:  # Don't make chunks too large
                        chunk_text = extended_text
                
                # Extract page number if available
                page_num = 1
                if '[Page ' in chunk_text:
                    try:
                        page_start = chunk_text.find('[Page ') + 6
                        page_end = chunk_text.find(']', page_start)
                        page_num = int(chunk_text[page_start:page_end])
                    except:
                        pass
                
                chunk = {
                    'chunk_id': f"chunk_{i}",
                    'text': chunk_text,
                    'source_url': source_url,
                    'page': page_num,
                    'chunk_index': i,
                    'length': len(chunk_text)
                }
                chunks.append(chunk)
            
            logger.info(f"Created {len(chunks)} chunks from {len(sentences)} sentences")
            
            if not chunks:
                raise ValueError("No valid chunks created from document")
                
            return chunks
            
        except Exception as e:
            logger.error(f"Error creating chunks: {e}")
            # Ultimate fallback to character-based chunking
            return self._fallback_character_chunks(text, source_url)
    
    def _simple_sentence_split(self, text: str) -> List[str]:
        """Simple sentence splitting fallback"""
        import re
        
        # Split on sentence endings
        sentences = re.split(r'[.!?]+\s+', text)
        
        # Clean up sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 20:  # Increased minimum sentence length
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def _fallback_character_chunks(self, text: str, source_url: str) -> List[Dict[str, Any]]:
        """Fallback to simple character-based chunking"""
        if self._is_text_garbage(text):
            raise ValueError("Cannot create chunks from corrupted text")
            
        chunk_size = 1500  # Larger chunks
        overlap = 300
        chunks = []
        
        for i in range(0, len(text), chunk_size - overlap):
            chunk_text = text[i:i + chunk_size]
            if len(chunk_text.strip()) < 100:
                continue
                
            chunks.append({
                'chunk_id': f"fallback_chunk_{i}",
                'text': chunk_text,
                'source_url': source_url,
                'page': 1,
                'chunk_index': i // chunk_size,
                'length': len(chunk_text)
            })
        
        logger.info(f"Created {len(chunks)} fallback chunks")
        
        if not chunks:
            raise ValueError("No valid chunks could be created")
            
        return chunks